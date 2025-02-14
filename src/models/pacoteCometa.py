from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
import os
import sys
import datetime

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def generate_cometa_contract(client, event, package, discount, folder):
    months = ['Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho', 'Julho', 'Agosto', 'Setembro', 'Outubro',
              'Novembro', 'Dezembro']
    date = datetime.datetime.now()

    #CONTRACT DETAILS
    extra_hour = "200,00"
    extra_hour_spelled = "duzentos reais"

    if event.payment_type == "PIX":
        coverage_price = package.pix_price - package.discounted_value(float(discount))
    else:
        coverage_price = package.card_price - package.discounted_value(float(discount))

    #OPENING NEW FILE
    document = Document()

    #EDITING GLOBAL STYLES
    style = document.styles['Normal']
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.color = "#000000"

    #EDITING FILE
    header = document.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run().add_picture(resource_path("src\\assets\\logo.png"), width=Pt(150), height=Pt(90))
    header.add_run().add_break(WD_BREAK.LINE)
    header.add_run('CONTRATO DE PRESTAÇÃO DE SERVIÇOS FOTOGRÁFICOS').bold = True

    personal_info = document.add_paragraph()
    personal_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    personal_info.add_run('Eveline Medeiros').bold = True
    personal_info.add_run(', fotógrafa, brasileira, portadora da carteira de identidade nº ')
    personal_info.add_run('6110281539').bold = True
    personal_info.add_run(', expedida pelo IFP, inscrita no CNPJ sob o número ')
    personal_info.add_run('32.973.993/0001-74').bold = True
    personal_info.add_run(', residente e domiciliada à ')
    personal_info.add_run('Rua Santo Antônio, 184, Vila City, Cachoeirinha, Rio Grande do Sul').bold = True
    personal_info.add_run(', doravante denominada ')
    personal_info.add_run('CONTRATADO').bold = True
    personal_info.add_run(', e ')
    personal_info.add_run(f'{client.name} ').bold = True
    personal_info.add_run('CPF: ')
    personal_info.add_run(f'{client.cpf}').bold = True
    personal_info.add_run(', residente em ')
    personal_info.add_run(f'{client.address}').bold = True
    personal_info.add_run(', doravante denominado(a) ')
    personal_info.add_run('CONTRATANTE').bold = True
    personal_info.add_run(', resolvem celebrar o presente contrato pelas cláusulas e condições a seguir:')

    clause_object_title = document.add_paragraph()
    clause_object_title.add_run('Cláusula Primeira - Objeto').bold = True
    clause_object = document.add_paragraph()
    clause_object.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    clause_object.add_run('O presente contrato destina-se à prestação de serviço(s) fotográfico(s) do(s) seguinte(s) evento(s): ')
    clause_object.add_run(f'{event.name} em {event.location} dia {event.date} à partir das {event.start_time} por {package.coverage_duration} de cobertura').bold = True
    clause_object.add_run('. No qual se obriga a entregar ao ')
    clause_object.add_run('CONTRATANTE').bold = True
    clause_object.add_run(', como objeto deste contrato, fotos do(s) evento(s), onde a quantidade delas é decidida unicamente pelo contratado (mínimo 100 fotos). As fotos devidamente selecionadas e editadas pelo ')
    clause_object.add_run('CONTRATADO').bold = True
    clause_object.add_run(f', serão entregues por email para {client.email} em até 30 dias após a data do evento. Após a seleção, o álbum e o estojo personalizado serão diagramados e um link será enviado para o ')
    clause_object.add_run('CONTRATANTE').bold = True
    clause_object.add_run(' para possíveis modificações ')
    clause_object.add_run('limite de até 2 modificações)')
    clause_object.add_run('. Após a autorização do ')
    clause_object.add_run('CONTRATANTE').bold = True
    clause_object.add_run(', o álbum e o estojo serão produzidos. O tempo de produção do fornecedor pode levar até 15 dias úteis, podendo esta estimativa sofrer alterações com base nos prazos do nosso fornecedor. Após este prazo, o álbum 25x25cm mais o estojo juntamente com o pendrive com a fotos digitais estarão disponíveis para retirada. ')
    clause_object.add_run('A entrega do kit completo fica à cargo do CONTRATANTE ').bold = True
    clause_object.add_run('podendo o mesmo retirar o kit em nosso estúdio na Av. Francisco Brochado da Rocha, 207 em horário agendado após a confecção do mesmo, ou solicitando algum serviço de entrega como o Uber Flash para o envio para a sua residência.')

    services = document.add_paragraph()
    services.add_run('Serviços adquiridos').bold = True
    services.add_run().add_break(WD_BREAK.LINE)
    services.add_run(f'Pacote {package.name}')

    clause_payment_title = document.add_paragraph()
    clause_payment_title.add_run('Cláusula Segunda - Valor e Forma de Pagamento').bold = True
    clause_payment = document.add_paragraph()
    clause_payment.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    clause_payment.add_run('O ')
    clause_payment.add_run('CONTRATANTE').bold = True
    clause_payment.add_run(', obriga-se a pagar ao ')
    clause_payment.add_run('CONTRATADO').bold = True
    clause_payment.add_run(', a importância total de ')
    clause_payment.add_run(f'R$ {'{0:.2f}'.format(coverage_price).replace(".", ",")} ').bold = True
    clause_payment.add_run('pela prestação dos serviços descritos na cláusula anterior mais ')
    clause_payment.add_run(f'R$ {'{0:.2f}'.format(float(event.commuting_fee)).replace(".", ",")}').bold = True
    clause_payment.add_run(' de deslocamento, totalizando ')
    clause_payment.add_run(f'R$ {'{0:.2f}'.format(coverage_price + float(event.commuting_fee)).replace(".", ",")}').bold = True
    clause_payment.add_run(', sendo este valor efetivado via ')
    clause_payment.add_run(f'{event.payment_type} até o dia {event.payment_due_date}').bold = True
    clause_payment.add_run('.')

    clause_payment_warning = document.add_paragraph()
    clause_payment_warning.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    clause_payment_warning.add_run('O atraso no pagamento gerará multa de 0,33% por dia atrasado ou 10% ao mês, podendo acarretar também na negativação do CPF nos serviços de proteção ao crédito.').bold = True
    clause_payment_warning.add_run().add_break(WD_BREAK.LINE)
    clause_payment_warning.add_run('Chave PIX: ').bold = True
    clause_payment_warning.add_run('amarinfancias@gmail.com')

    unique_paragraph = document.add_paragraph()
    unique_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    unique_paragraph.add_run('Parágrafo Único').bold = True
    unique_paragraph.add_run(' - Caso o ')
    unique_paragraph.add_run('CONTRATANTE ').bold = True
    unique_paragraph.add_run('tenha interesse de acrescentar hora a mais de serviço prestado no evento, as mesmas serão cobradas a parte pelo ')
    unique_paragraph.add_run('CONTRATADO ').bold = True
    unique_paragraph.add_run(f'ao valor de R$ {extra_hour} ({extra_hour_spelled}) por hora extra.')

    obligations_3_title = document.add_paragraph()
    obligations_3_title.add_run('Cláusula Terceira - Obrigações do Contratado').bold = True
    obligations_3 = document.add_paragraph()
    obligations_3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_3.add_run('3. Caso haja alguma falha no equipamento do ')
    obligations_3.add_run('CONTRATADO').bold = True
    obligations_3.add_run(', alheia à sua vontade e só perceptível após a execução do serviço, que acarrete perda total ou parcial do serviço prestado, fica o ')
    obligations_3.add_run('CONTRATADO').bold = True
    obligations_3.add_run(' obrigado a devolver a quantia proporcional à perda ocorrida, até o limite do valor contratado, ou poderá compensar o ')
    obligations_3.add_run('CONTRATANTE ').bold = True
    obligations_3.add_run('de outra maneira, na concordância das partes.')

    obligations_3_1 = document.add_paragraph()
    obligations_3_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_3_1.add_run('3.1 O ')
    obligations_3_1.add_run('CONTRATADO ').bold = True
    obligations_3_1.add_run('não se responsabilizará por problemas causados por fenômenos da natureza e/ou terceiros que, de qualquer modo, venham a prejudicar o bom andamento do serviço contratado.')

    obligations_3_2 = document.add_paragraph()
    obligations_3_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_3_2.add_run('3.2 O ')
    obligations_3_2.add_run('CONTRATADO ').bold = True
    obligations_3_2.add_run('fica isento de quaisquer responsabilidades caso não consiga exercer seu trabalho com seus equipamentos habituais, em virtude de proibições que venham a ocorrer por parte da administração do(s) local(is) onde ocorrerá(ão) o(s) evento(s).')

    obligations_3_3 = document.add_paragraph()
    obligations_3_3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_3_3.add_run('3.3 Fica acertado que, em caso de motivo de força maior justificável que impossibilite a pessoa do ')
    obligations_3_3.add_run('CONTRATADO ').bold = True
    obligations_3_3.add_run('estar presente na data do trabalho para a sua cobertura fotográfica, o ')
    obligations_3_3.add_run('CONTRATANTE ').bold = True
    obligations_3_3.add_run('desde já autoriza o ')
    obligations_3_3.add_run('CONTRATADO ').bold = True
    obligations_3_3.add_run('a substituir sua presença por outro profissional do mesmo nível técnico e portando equipamentos semelhantes.')

    obligations_4_title = document.add_paragraph()
    obligations_4_title.add_run('Cláusula Quarta – Obrigações do Contratante').bold = True
    obligations_4 = document.add_paragraph()
    obligations_4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_4.add_run('4. O ')
    obligations_4.add_run('CONTRATANTE ').bold = True
    obligations_4.add_run('deverá fornecer ao ')
    obligations_4.add_run('CONTRATADO ').bold = True
    obligations_4.add_run('todas as informações necessárias à realização do serviço, devendo especificar os detalhes necessários à perfeita consecução do mesmo; possibilitar o livre acesso da equipe ao local do evento; verificar a existência de pontos de energia para os equipamentos (iluminação, câmeras, etc); reservar uma mesa para a equipe se instalar quando necessário; fornecer alimentação à equipe durante o evento.')

    obligations_4_1 = document.add_paragraph()
    obligations_4_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_4_1.add_run('4.1 É de responsabilidade do ')
    obligations_4_1.add_run('CONTRATADO ').bold = True
    obligations_4_1.add_run('por parte de seus convidados e/ou por parte da equipe de apoio do evento (ex: segurança, cerimonial, garçom, etc.) devendo ressarcir o prejuízo causado ao ')
    obligations_4_1.add_run('CONTRATADO').bold = True
    obligations_4_1.add_run('.')

    obligations_4_2 = document.add_paragraph()
    obligations_4_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    obligations_4_2.add_run('4.2 As despesas adicionais, tais como estacionamento, taxa de entrada, ingressos, taxas de igrejas, cartórios ou locais do evento e locações serão de única responsabilidade do ')
    obligations_4_2.add_run('CONTRATANTE').bold = True
    obligations_4_2.add_run('.')

    general_clause_5_title = document.add_paragraph()
    general_clause_5_title.add_run('Cláusula Quinta – Disposições Gerais').bold = True
    general_clause_5 = document.add_paragraph()
    general_clause_5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5.add_run('5. Desde já, os negativos provenientes deste contrato de prestação de serviço, são de propriedade do ')
    general_clause_5.add_run('CONTRATADO').bold = True
    general_clause_5.add_run(', cabendo ao mesmo os Créditos e Direitos Autorais, conforme Lei 9.610, de 20/02/98. Havendo interesse do ')
    general_clause_5.add_run('CONTRATANTE ').bold = True
    general_clause_5.add_run('em adquiri-los, os mesmos serão negociados à parte deste contrato.')

    general_clause_5_1 = document.add_paragraph()
    general_clause_5_1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_1.add_run('5.1 As fotografias pertencem ao ')
    general_clause_5_1.add_run('CONTRATADO ').bold = True
    general_clause_5_1.add_run('pela Lei 9.610, de 20/02/98 (autor da obra) e ficarão sob guarda deste. Portanto é proibida a venda de alguma dessas fotografias sem o conhecimento do ')
    general_clause_5_1.add_run('CONTRATADO ').bold = True
    general_clause_5_1.add_run(', e tendo assim seus devidos lucros. É também proibida a publicação dessas fotografias em sites, jornais, revistas sociais e comerciais sem o conhecimento e devidos direitos autorais do ')
    general_clause_5_1.add_run('CONTRATADO').bold = True
    general_clause_5_1.add_run('. Toda a ação que gerar lucro em cima de qualquer obra do ')
    general_clause_5_1.add_run('CONTRATADO ').bold = True
    general_clause_5_1.add_run('deverá ser acordada através de uma negociação, e feito contrato sobre os direitos do ')
    general_clause_5_1.add_run('CONTRATADO ').bold = True
    general_clause_5_1.add_run(' sobre esta imagem(ns) e o seus lucros.')

    general_clause_5_2 = document.add_paragraph()
    general_clause_5_2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_2.add_run('5.2 Tendo por assinar este contrato o ')
    general_clause_5_2.add_run('CONTRATANTE ').bold = True
    general_clause_5_2.add_run('está autorizando o ')
    general_clause_5_2.add_run('CONTRATADO ').bold = True
    general_clause_5_2.add_run('a usar as imagens tiradas neste serviço em seu portfólio, site, blog, facebook, instagram e outras ')
    general_clause_5_2.add_run('redes sociais e físicas onde divulga o seu trabalho').bold = True
    general_clause_5_2.add_run('.')

    general_clause_5_3 = document.add_paragraph()
    general_clause_5_3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_3.add_run('5.3 Tendo como confirmação na assinatura deste contrato, o ')
    general_clause_5_3.add_run('CONTRATANTE ').bold = True
    general_clause_5_3.add_run('concorda que após a entrega do objeto deste contrato descrito na cláusula primeira, o ')
    general_clause_5_3.add_run('CONTRATADO ').bold = True
    general_clause_5_3.add_run('não tem como responsabilidade e obrigação guardar as fotos originais e editadas deste serviço ')
    general_clause_5_3.add_run('por mais de 1 mês após a entrega do mesmo').bold = True
    general_clause_5_3.add_run('.')

    general_clause_5_4 = document.add_paragraph()
    general_clause_5_4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_4.add_run('5.4 O ')
    general_clause_5_4.add_run('CONTRATANTE ').bold = True
    general_clause_5_4.add_run('confirma neste contrato estar ciente ')
    general_clause_5_4.add_run('de que não serão feitas correções e manipulações digitais no Photoshop, e programas similares, tais como: correção de pele, alteração de corpo, retirada de fundo da imagem, entre outras, em nenhuma das imagens').bold = True
    general_clause_5_4.add_run('. Caso o ')
    general_clause_5_4.add_run('CONTRATANTE ').bold = True
    general_clause_5_4.add_run('solicite alguma dessas alterações será feito um contrato à parte com a devida negociação de valores, pois, neste contrato, este serviço não está incluso.')

    general_clause_5_5 = document.add_paragraph()
    general_clause_5_5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_5.add_run('5.5 O ')
    general_clause_5_5.add_run('CONTRATANTE ').bold = True
    general_clause_5_5.add_run('receberá prévias das fotos decorrentes do serviço contratado ficando o mesmo responsável por avisar de quaisquer correções desejadas, desde que ')
    general_clause_5_5.add_run('não estejam em conflito com as outras cláusulas do contrato, antes do envio do material completo')
    general_clause_5_5.add_run('.')

    general_clause_5_6 = document.add_paragraph()
    general_clause_5_6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_6.add_run('5.6 O ')
    general_clause_5_6.add_run('CONTRATANTE ').bold = True
    general_clause_5_6.add_run('confirma neste contrato estar ciente de que ')
    general_clause_5_6.add_run('não estão autorizadas manipulações das imagens').bold = True
    general_clause_5_6.add_run(', tais como: ')
    general_clause_5_6.add_run('filtros de cor, corte, correções de luz, manipulação de fundo').bold = True
    general_clause_5_6.add_run(', entre outras, do objeto deste contrato após a sua entrega, garantindo assim a preservação da identidade visual e pós-produção do ')
    general_clause_5_6.add_run('CONTRATADO').bold = True
    general_clause_5_6.add_run('.')

    general_clause_5_7 = document.add_paragraph()
    general_clause_5_7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_7.add_run('5.7 O ')
    general_clause_5_7.add_run('CONTRATANTE ').bold = True
    general_clause_5_7.add_run('fica ciente de que as fotos serão editadas utilizando os presets/edições padrão da empresa AMAR Infâncias, não sendo possível realizar a entrega do material sem a edição padrão e/ou com preset diferente dos padrões da empresa.')

    general_clause_5_8 = document.add_paragraph()
    general_clause_5_8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_8.add_run('5.8 O ')
    general_clause_5_8.add_run('CONTRATANTE ').bold = True
    general_clause_5_8.add_run('confirma estar ciente de que quaisquer alterações solicitadas após a entrega de produtos que necessitam de revelação, terão os custos de reimpressões, revelações adicionais ou quaisquer custos relacionados a novas confecções destes materiais ')
    general_clause_5_8.add_run('revertidos em 100% para o CONTRATANTE').bold = True
    general_clause_5_8.add_run('.')

    general_clause_5_9 = document.add_paragraph()
    general_clause_5_9.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_9.add_run('5.9 O ')
    general_clause_5_9.add_run('CONTRATANTE').bold = True
    general_clause_5_9.add_run(' confirma neste contrato estar ciente de que na entrega de pacotes com edição de vídeo será aceita e realizada apenas 1 (uma) alteração na edição.')

    general_clause_5_10 = document.add_paragraph()
    general_clause_5_10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    general_clause_5_10.add_run('5.10 Fica compactuado entre as partes a total inexistência de vínculo trabalhista entre ambas, excluindo as obrigações previdenciárias e os encargos sociais, não havendo entre ')
    general_clause_5_10.add_run('CONTRATADO ').bold = True
    general_clause_5_10.add_run('e ')
    general_clause_5_10.add_run('CONTRATANTE ').bold = True
    general_clause_5_10.add_run('qualquer tipo de relação de subordinação.')

    breach_of_contract_title = document.add_paragraph()
    breach_of_contract_title.add_run('Cláusula sexta - Quebra de Contrato').bold = True
    breach_of_contract = document.add_paragraph()
    breach_of_contract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    breach_of_contract.add_run('Em caso de quebra de contrato por parte do ')
    breach_of_contract.add_run('CONTRATANTE').bold = True
    breach_of_contract.add_run(' seja por necessidade de cancelamento do serviço agendado ou por quaisquer outras alterações nos dispostos deste documento que acarretem na impossibilidade da celebração do contrato da forma acordada por ambas as partes, o ')
    breach_of_contract.add_run('CONTRATADO ').bold = True
    breach_of_contract.add_run('deve reembolsar o valor dos serviços prestados em 60% do valor total acordado, caso o pagamento já tenha sido efetuado. Em caso de pagamento de sinal, reserva de horário ou qualquer outro tipo de pagamento similar, estes valores não serão reembolsados.')

    contract_duration_title = document.add_paragraph()
    contract_duration_title.add_run('Cláusula sétima - Da Vigência do Contrato').bold = True
    contract_duration = document.add_paragraph()
    contract_duration.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    contract_duration.add_run('Caso o ')
    contract_duration.add_run('CONTRATANTE').bold = True
    contract_duration.add_run(' não realize a assinatura do presente contrato, o pagamento pelos serviços e/ou a realização do serviço fotográfico equivalem a aceite automático de todas as cláusulas constantes no presente documento.')

    document.add_paragraph('E, por estarem de acordo, confirmam este contrato:')
    document.add_paragraph(f'Cachoeirinha, {months[int(date.strftime('%m')) - 1]} de {date.strftime("%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER

    client_signature = document.add_paragraph()
    client_signature.add_run('CONTRATANTE:')
    client_signature.add_run().add_break(WD_BREAK.LINE)
    client_signature.add_run(f'{client.name}')
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph()
    document.add_paragraph('________________________________________________')

    amar_signature = document.add_paragraph()
    amar_signature.add_run('CONTRATADA:')
    amar_signature.add_run().add_break(WD_BREAK.LINE)
    amar_signature.add_run('Eveline Medeiros')
    signature = document.add_paragraph()
    signature.add_run().add_picture(resource_path('src\\assets\\signature.png'), width=Pt(80), height=Pt(80))
    document.add_paragraph('________________________________________________')

    #SAVING
    document.save(f'{folder}/Contrato - {client.name}.docx')